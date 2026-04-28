from docx import Document
from .supabase_client import get_chapters
import os
from datetime import datetime


def slugify(text: str) -> str:
    return ''.join(c if c.isalnum() else '_' for c in text).strip('_')[:100]


def compile_book_to_docx(book_request: dict, out_dir='outputs'):
    os.makedirs(out_dir, exist_ok=True)
    book_id = book_request.get('id')
    title = book_request.get('title')
    filename = f"{book_id}_{slugify(title)}.docx"
    path = os.path.join(out_dir, filename)

    doc = Document()
    # Title
    doc.add_heading(title or f'Book {book_id}', level=0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")
    doc.add_page_break()

    # Outline (if present)
    outline = book_request.get('outline')
    if outline:
        doc.add_heading('Outline', level=1)
        if isinstance(outline, dict) and 'chapters' in outline:
            for i, ch in enumerate(outline['chapters'], start=1):
                doc.add_heading(f"{i}. {ch.get('title', '')}", level=2)
                bullets = ch.get('bullets') or ch.get('points') or []
                for b in bullets:
                    doc.add_paragraph(b, style='List Bullet')
        else:
            doc.add_paragraph(str(outline))
        doc.add_page_break()

    # Chapters
    chapters = get_chapters(book_id) or []
    for ch in chapters:
        num = ch.get('chapter_number')
        doc.add_heading(f"Chapter {num}: {ch.get('chapter_title') or ''}", level=1)
        content = ch.get('chapter_content') or ''
        # naive split into paragraphs
        for p in content.split('\n\n'):
            doc.add_paragraph(p)
        doc.add_page_break()

    doc.save(path)
    return path
